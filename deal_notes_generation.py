import os
import json
from dotenv import load_dotenv
from google import genai
from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    return content

def save_to_docx(text, filename):
    doc = Document()
    doc.add_heading("Deal Notes Generation", 0)
    if text is None:
        doc.add_paragraph("No deal notes were generated. Please check your input or model response.")
    else:
        for paragraph in text.split('\n\n'):
            doc.add_paragraph(paragraph)
    doc.save(filename)

def read_data(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data

def generate_deal_notes(deal_notes_prompt, startup_information, client):
    full_input = (
        f"{deal_notes_prompt}\n\n"
        "Here is the extracted startup data in JSON format:\n"
        f"{json.dumps(startup_information, indent=2)}"
    )
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[
            {
                "role": "user",
                "parts": [{"text": full_input}]
            }
        ]
    )
    # Defensive: check that response.text is not None
    return getattr(response, "text", None)

def display_deal_notes(notes_text, st):
    st.markdown("""
        <style>
        .section {
            background: #eaf6ff; border-radius: 10px; padding: 12px; margin-bottom: 12px; color: #0a2540; font-size: 16px;
        }
        .summary {
            background: #f7e7ff; border-radius: 10px; padding: 12px; margin-bottom: 12px; color: #8e24aa; font-size: 18px; font-weight: bold;
        }
        .strength {
            background: #d0ffd6; border-radius: 10px; padding: 12px; margin-bottom: 12px; color: #388e3c; font-size: 16px;
        }
        .risk {
            background: #ffd6d6; border-radius: 10px; padding: 12px; margin-bottom: 12px; color: #d32f2f; font-size: 16px;
        }
        .recommendation {
            background: #fff7d6; border-radius: 10px; padding: 12px; margin-bottom: 12px; color: #ff8f00; font-size: 16px;
        }
        </style>
    """, unsafe_allow_html=True)

    if not notes_text:
        st.warning("No deal notes to display. Please generate deal notes first.")
        return

    sections = notes_text.split('\n\n')
    for section in sections:
        section_lower = section.lower()
        if 'summary' in section_lower:
            st.markdown(f'<div class="summary">📝 <b>Summary:</b><br>{section}</div>', unsafe_allow_html=True)
        elif 'strength' in section_lower or 'positive' in section_lower:
            st.markdown(f'<div class="strength">💪 <b>Strengths:</b><br>{section}</div>', unsafe_allow_html=True)
        elif 'risk' in section_lower or 'concern' in section_lower:
            st.markdown(f'<div class="risk">⚠️ <b>Risks:</b><br>{section}</div>', unsafe_allow_html=True)
        elif 'recommendation' in section_lower or 'suggestion' in section_lower:
            st.markdown(f'<div class="recommendation">🔖 <b>Recommendations:</b><br>{section}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="section">{section}</div>', unsafe_allow_html=True)